from docx import Document

def text_to_binary(text):
    binary_text = ''.join(format(ord(char), '08b') for char in text)
    return binary_text

def word_to_binary(word):
    binary_word = ''.join(format(ord(char), '08b') for char in word)
    return binary_word

def modify_spaces(word, bit):
    if bit == '0':
        return word + ""
    elif bit == '1':
        return word + " "
    else:
        raise ValueError("Некорректный бит")

def steganography(input_file, message, output_file):
    doc = Document(input_file)
    modified_doc = Document()
    binary_message = text_to_binary(message)
    message_index = 0

    for paragraph in doc.paragraphs:
        modified_paragraph = modified_doc.add_paragraph()
        for run in paragraph.runs:
            words = run.text.split()
            for i, word in enumerate(words):
                if word.strip():
                    if binary_message and message_index < len(binary_message):
                        bit = binary_message[message_index]
                        modified_word = modify_spaces(word, bit)
                        binary_message = binary_message[1:]
                        message_index += 1
                    else:
                        modified_word = word
                    modified_paragraph.add_run(modified_word)
                else:
                    modified_paragraph.add_run(" ")
                if i < len(words) - 1:
                    modified_paragraph.add_run(" ")

    modified_doc.save(output_file)

if __name__ == "__main__":
    input_file = "input.docx"
    message = "Kirill"
    output_file = "output.docx"
    steganography(input_file, message, output_file)
    print("Стеганография завершена. Результат сохранен в", output_file)
