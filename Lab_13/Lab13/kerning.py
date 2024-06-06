from docx import Document

def apply_kerning(text, bit):
    kerning_pairs = ['th', 'he', 'an', 'nd', 'di', 'ia', 'ne', 'ed', 'al', 'is', 'la', 'in', 'to', 'se', 'as', 'fo', 'me', 'il', 'of']

    letters = list(text)
    kerned_pairs = []

    for i in range(len(letters) - 1):
        pair = letters[i] + letters[i + 1]
        if pair in kerning_pairs and bit == '1':
            letters[i] += ' '
            kerned_pairs.append(pair)

    kerned_text = ''.join(letters)

    return kerned_text, kerned_pairs

def steganography(input_file, message, output_file):
    doc = Document(input_file)
    modified_doc = Document()
    binary_message = ''.join(format(ord(char), '08b') for char in message)
    message_index = 0

    for paragraph in doc.paragraphs:
        modified_paragraph = modified_doc.add_paragraph()
        for run in paragraph.runs:
            words = run.text.split()
            for word in words:
                if word.strip():
                    if binary_message and message_index < len(binary_message):
                        bit = binary_message[message_index]
                        modified_word, kerned_pairs = apply_kerning(word, bit)
                        binary_message = binary_message[1:]
                        message_index += 1
                    else:
                        modified_word = word
                    modified_paragraph.add_run(modified_word)
                    if binary_message and words.index(word) < len(words) - 1:
                        modified_paragraph.add_run(" ")
                else:
                    modified_paragraph.add_run(" ")

    modified_doc.save(output_file)
    print("Стеганография завершена. Результат сохранен в", output_file)

if __name__ == "__main__":
    input_file = "input.docx"
    message = "Kirill"
    output_file = "output.docx"
    steganography(input_file, message, output_file)

# from docx import Document

# def text_to_binary(text):
#     binary_text = ''.join(format(ord(char), '08b') for char in text)
#     return binary_text

# def word_to_binary(word):
#     binary_word = ''.join(format(ord(char), '08b') for char in word)
#     return binary_word

# def modify_kerning(word, bit):
#     if bit == '0':
#         return "".join([char + "" for char in word])  # Значение кернинга 0px для слов, начинающихся с нулевого бита
#     elif bit == '1':
#         return "".join([char + " " for char in word])  # Значение кернинга 1px для слов, начинающихся с единичного бита
#     else:
#         raise ValueError("Некорректный бит")

# def steganography(input_file, message, output_file):
#     doc = Document(input_file)
#     modified_doc = Document()
#     binary_message = text_to_binary(message)
#     message_index = 0

#     for paragraph in doc.paragraphs:
#         modified_paragraph = modified_doc.add_paragraph()
#         for run in paragraph.runs:
#             words = run.text.split()
#             for word in words:
#                 if word.strip():
#                     if binary_message and message_index < len(binary_message):
#                         word_binary = word_to_binary(word)
#                         bit = binary_message[message_index]
#                         modified_word = modify_kerning(word, bit)
#                         binary_message = binary_message[1:]
#                         message_index += 1
#                     else:
#                         modified_word = word
#                     modified_paragraph.add_run(modified_word)
#                     # Добавляем пробел между словами, если осталось еще сообщение для скрытия
#                     if binary_message and words.index(word) < len(words) - 1:
#                         modified_paragraph.add_run(" ")
#                 else:
#                     modified_paragraph.add_run(" ")  # Добавляем пробел между словами

#     modified_doc.save(output_file)

# if __name__ == "__main__":
#     input_file = "input.docx"
#     message = "Kirill"
#     output_file = "output.docx"
#     steganography(input_file, message, output_file)
#     print("Стеганография завершена. Результат сохранен в", output_file)
